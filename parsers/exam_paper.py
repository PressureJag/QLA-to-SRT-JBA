"""
Parses maths exam paper Word documents (.doc / .docx) to extract a question bank.

Questions are keyed by their number string with a group + paper prefix:
  "CP_1", "CP_2" …        — Core Plus Paper 1
  "CP_P2_1", "CP_P2_2" … — Core Plus Paper 2
  "N_1", "N_2" …          — Nurture/Numeracy Paper 1
  "N_P2_1" …              — Nurture/Numeracy Paper 2

If the group cannot be detected from the filename, the prefix is omitted
and the old unprefixed keys ("1", "P2_1") are used as a fallback.

The QLA "Question number" row uses the same numbers (and sub-part suffixes
like "6a", "6b") — analyser.py resolves them at lookup time.
"""

import os
import re
import subprocess
import tempfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None

# Matches question lines like "1\tWrite 180 minutes..." or "10\tAB and BC..."
_Q_START     = re.compile(r"^(\d+)\t")
# Prefers the definitive mark total line; falls back to bare "(N marks)"
_TOTAL_MARKS = re.compile(r"\(Total for Question \d+[a-z]* is (\d+) marks?\)", re.IGNORECASE)
_MARKS       = re.compile(r"\((\d+)\s*marks?\)", re.IGNORECASE)

import re as _re

# Ordered patterns: first match wins.
# Use (?<![a-zA-Z]) / (?![a-zA-Z]) instead of \b so that underscores
# (inserted by werkzeug's secure_filename) act as separators.
_GROUP_PATTERNS = [
    # Full words / phrases — most specific first
    (_re.compile(r'core[\s_-]*plus',                   _re.I), "CP"),
    (_re.compile(r'(?<![a-zA-Z])nurture(?![a-zA-Z])',  _re.I), "N"),
    (_re.compile(r'(?<![a-zA-Z])numeracy(?![a-zA-Z])', _re.I), "N"),
    (_re.compile(r'(?<![a-zA-Z])extension(?![a-zA-Z])',_re.I), "EXTENSION"),
    # Short codes (uppercase only to avoid matching inside other words)
    (_re.compile(r'(?<![A-Za-z])N(?![A-Za-z])'),              "N"),
    (_re.compile(r'(?<![A-Za-z])EXT(?![A-Za-z])'),            "EXTENSION"),
    (_re.compile(r'(?<![A-Za-z])CP(?![A-Za-z])'),             "CP"),
]


def _detect_group_paper(path):
    """
    Guess the ability group and paper number from the filename.
    Returns (group_code_or_None, 1_or_2).
    """
    stem = Path(str(path)).stem          # without extension
    name_lower = stem.lower()
    group = None
    for pattern, code in _GROUP_PATTERNS:
        if pattern.search(stem):         # search on original (preserves case for short codes)
            group = code
            break
    paper = 2 if ("paper 2" in name_lower or "paper2" in name_lower or "paper_2" in name_lower) else 1
    return group, paper


def _to_docx(path):
    """
    Return a path to a .docx version of the file.
    If the file is already .docx, returns (original_path, False).
    If .doc, converts via macOS textutil and returns (tmp_path, True).
    Caller is responsible for deleting the tmp file when is_temp=True.
    """
    if Path(str(path)).suffix.lower() == ".docx":
        return str(path), False
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        subprocess.run(
            ["textutil", "-convert", "docx", "-output", tmp.name, str(path)],
            check=True, capture_output=True, timeout=30,
        )
        return tmp.name, True
    except Exception:
        os.unlink(tmp.name)
        raise


def _extract_questions(source_path):
    """
    Walk every paragraph in a Word document and group text under Q-number headers.
    Accepts both .docx and .doc (the latter is converted via macOS textutil).
    Returns {"1": {text, marks}, "2": {text, marks}, ...}
    """
    if Document is None:
        raise ImportError("python-docx is required: pip install python-docx")

    docx_path, is_temp = _to_docx(source_path)
    try:
        doc = Document(str(docx_path))
        bank = {}
        current_key = None
        current_parts = []

        def _flush():
            if not current_key or not current_parts:
                return
            full = " ".join(p.strip() for p in current_parts if p.strip())
            m = _TOTAL_MARKS.search(full) or _MARKS.search(full)
            bank[current_key] = {
                "text":  full,
                "marks": int(m.group(1)) if m else None,
            }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            m = _Q_START.match(text)
            if m:
                _flush()
                current_key = m.group(1)   # "1", "2", "10", "23" etc.
                current_parts = [text]
            elif current_key:
                current_parts.append(text)

        _flush()

        # Also scan table cells — some papers embed questions in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue
                        m = _Q_START.match(text)
                        if m:
                            key = m.group(1)
                            if key not in bank:
                                bank[key] = {"text": text, "marks": None}
                                m2 = _TOTAL_MARKS.search(text) or _MARKS.search(text)
                                if m2:
                                    bank[key]["marks"] = int(m2.group(1))

        return bank
    finally:
        if is_temp and os.path.exists(docx_path):
            os.unlink(docx_path)


def parse_papers(paper_paths):
    """
    Parse one or more exam paper Word documents.

    paper_paths — list of Path/str objects (None entries are skipped).

    The group and paper number are detected from each filename, so you can
    pass papers for multiple ability groups at once:
      - "Year 7 - Core Plus - ... - Paper 1.doc" → keys "CP_1", "CP_2" ...
      - "Year 7 - Core Plus - ... - Paper 2.doc" → keys "CP_P2_1" ...
      - "Year 7 - Nurture  - ... - Paper 1.doc"  → keys "N_1", "N_2" ...

    If the group cannot be detected from the filename, plain keys are used
    as a fallback ("1", "P2_1") for backwards compatibility.

    Returns a merged question bank dict, or {} if no valid paths supplied.
    """
    bank = {}

    for path in (paper_paths or []):
        if not path or not Path(str(path)).exists():
            continue
        group, paper_num = _detect_group_paper(path)
        try:
            extracted = _extract_questions(path)
        except Exception:
            continue
        p2 = "P2_" if paper_num == 2 else ""
        for q_num, val in extracted.items():
            raw_key = f"{p2}{q_num}"                          # e.g. "P2_5"
            key     = f"{group}_{raw_key}" if group else raw_key  # e.g. "CP_P2_5"
            bank[key] = val

    return bank
