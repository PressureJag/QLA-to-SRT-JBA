"""
SRT Question Sheet Generator
1 targeted question per weak topic per class → Word .docx
"""

import datetime
import os
import re
import zipfile
from io import BytesIO
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

SCHOOL_NAME = "Outwood Academy Newbold"
OUTPUT_DIR  = Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

FORMAT_GUIDE = {
    "short_answer": (
        "Write 1 short-answer question worth 1–3 marks. "
        "Include answer line(s): .......................................................... "
        "End with: (Total for this question: X mark[s])"
    ),
    "multiple_choice": (
        "Write 1 multiple-choice question. "
        "Give exactly 4 options: A) B) C) D). "
        "State the correct answer on its own line: Answer: [letter]"
    ),
    "exam_style": (
        "Write 1 exam-style question (may have parts (a)(b)). "
        "Include mark allocations in brackets. "
        "End with: (Total for this question: X mark[s])"
    ),
}


def _client():
    return anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))


def generate_class_questions(class_name, group_name, ability, fmt, weak_topics, year_group):
    """Call Claude to produce 1 question per weak topic for a single class."""
    topics_lines = "\n".join(
        f"{i+1}. {t['topic']} — class avg {t['avg_pct']:.1f}%  "
        f"({int(t['max_marks'])} mark{'s' if t['max_marks'] != 1 else ''})"
        for i, t in enumerate(weak_topics)
    )

    prompt = f"""You are an expert secondary maths teacher at {SCHOOL_NAME} writing SRT (Student Response Time) intervention questions.

Year Group: {year_group}
Class: {class_name}
Ability Group: {group_name} ({ability})

The 5 weakest topics for this class (worst first):
{topics_lines}

Task: write EXACTLY 1 question for each topic.

Question style: {FORMAT_GUIDE.get(fmt, FORMAT_GUIDE['short_answer'])}

Rules:
- Each question must test the exact skill named in the topic
- Match difficulty to {ability} ability Year {year_group}
- Topics below 35% average: use a scaffolded, accessible approach
- No answers shown (except Answer: line for multiple choice)
- No preamble — output only the 5 topic blocks

Use this structure exactly for each block:
TOPIC: [topic name]
CLASS_AVG: [x.x]%
Q1. [question text]
[answer lines / MC options / working space]
(Total for this question: X mark[s])
___
"""

    msg = _client().messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    return msg.content[0].text


# ── Word document builder ─────────────────────────────────────────────────────

def _run(para, text, bold=False, italic=False, size=None, rgb=None):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    if size:  r.font.size = Pt(size)
    if rgb:   r.font.color.rgb = RGBColor(*rgb)
    return r


def _render_questions(doc, text, fmt):
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            doc.add_paragraph()
            continue

        if s.startswith("TOPIC:"):
            p = doc.add_paragraph()
            _run(p, s.replace("TOPIC:", "TOPIC: ").strip(), bold=True, size=13, rgb=(0x1F, 0x49, 0x7D))

        elif s.startswith("CLASS_AVG:"):
            p = doc.add_paragraph()
            label = "Class average on this topic: " + s.split(":", 1)[1].strip()
            _run(p, label, italic=True, size=9, rgb=(0xBB, 0x33, 0x33))

        elif re.match(r"^Q\d+[a-z]?\.", s):
            doc.add_paragraph()
            p = doc.add_paragraph()
            _run(p, s, bold=True, size=11)

        elif re.match(r"^[A-D]\)", s):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            _run(p, s, size=11)

        elif s.lower().startswith("answer:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            _run(p, s, size=10, rgb=(0x00, 0x80, 0x00))

        elif "......" in s:
            p = doc.add_paragraph()
            _run(p, s, size=11)

        elif s.startswith("(Total for"):
            p = doc.add_paragraph()
            _run(p, s, italic=True, size=10)

        elif s.startswith("___"):
            p = doc.add_paragraph()
            _run(p, "─" * 95, size=8, rgb=(0xCC, 0xCC, 0xCC))
            doc.add_paragraph()

        elif re.match(r"^\([a-z]\)", s):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            _run(p, s, bold=True, size=11)

        else:
            p = doc.add_paragraph()
            _run(p, s, size=11)

    if fmt == "exam_style":
        doc.add_paragraph()
        tbl = doc.add_table(rows=6, cols=1)
        tbl.style = "Table Grid"
        for row in tbl.rows:
            row.height = Cm(0.9)


def build_class_docx(class_name, group_name, ability, fmt, year_group, overall_avg, weak_topics, questions_text):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    date_str = datetime.datetime.now().strftime("%d %B %Y")

    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(h, SCHOOL_NAME, bold=True, size=16, rgb=(0x1F, 0x49, 0x7D))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sub, f"Year {year_group}  |  {group_name}  |  Maths — SRT Intervention Questions",
         bold=True, size=12)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    avg_str = f"{overall_avg:.1f}%" if overall_avg is not None else "N/A"
    _run(info, f"Class: {class_name}  |  Ability: {ability}  |  Class Average: {avg_str}  |  {date_str}",
         italic=True, size=10, rgb=(0x55, 0x55, 0x55))

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(rule, "─" * 95, size=8, rgb=(0xAA, 0xAA, 0xAA))

    # Topic summary box
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "Topics addressed in this sheet:", bold=True, size=9, rgb=(0x55, 0x55, 0x55))
    for i, t in enumerate(weak_topics, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        _run(p, f"{i}. {t['topic']} ({t['avg_pct']:.1f}%)", size=9, rgb=(0x77, 0x77, 0x77))
    doc.add_paragraph()

    _render_questions(doc, questions_text, fmt)
    return doc


# ── Orchestrator ──────────────────────────────────────────────────────────────

def generate_srt_sheets(analysis, selected_groups=None):
    """
    Generate one Word .docx per class for the selected groups.
    Returns list of file-info dicts.
    """
    year_group    = analysis["year_group"]
    date_safe     = datetime.datetime.now().strftime("%Y-%m-%d")
    groups_to_run = selected_groups or list(analysis["groups"].keys())
    saved         = []

    for code in groups_to_run:
        group = analysis["groups"].get(code)
        if not group or not group.get("has_scores"):
            continue

        for class_name, cls_data in group["classes"].items():
            # topics is sorted weakest first by the analyser
            weak = cls_data.get("topics", [])[:5]
            if not weak:
                continue

            questions_text = generate_class_questions(
                class_name=class_name,
                group_name=group["name"],
                ability=group["ability"],
                fmt=group["format"],
                weak_topics=weak,
                year_group=year_group,
            )

            doc = build_class_docx(
                class_name=class_name,
                group_name=group["name"],
                ability=group["ability"],
                fmt=group["format"],
                year_group=year_group,
                overall_avg=cls_data.get("overall_avg_pct"),
                weak_topics=weak,
                questions_text=questions_text,
            )

            fname = f"SRT_{class_name}_Year{year_group}_{date_safe}.docx"
            fpath = OUTPUT_DIR / fname
            doc.save(fpath)
            saved.append({
                "class_name": class_name,
                "group_name": group["name"],
                "group_code": code,
                "ability":    group["ability"],
                "avg":        cls_data.get("overall_avg_pct"),
                "filename":   fname,
                "path":       str(fpath),
            })

    return saved


def create_zip(file_list, zip_name="SRT_Intervention_Sheets.zip"):
    """Bundle all .docx files into a ZIP and return a BytesIO buffer."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in file_list:
            zf.write(item["path"], item["filename"])
    buf.seek(0)
    return buf
